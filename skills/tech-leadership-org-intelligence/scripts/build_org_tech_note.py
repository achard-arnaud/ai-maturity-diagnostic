#!/usr/bin/env python3
"""Build a traceable technology-leadership organisation note from JSON."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - environment dependent
    raise SystemExit("Missing dependency: install python-docx") from exc

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError as exc:  # pragma: no cover - environment dependent
    raise SystemExit("Missing dependency: install pillow") from exc


GREEN = "4D8F2A"
DARK_GREEN = "2D5E20"
NAVY = "18324B"
TEAL = "1C6A73"
LIGHT_GREEN = "EAF3E5"
LIGHT_BLUE = "EAF1F7"
LIGHT_GREY = "F2F4F6"
MID_GREY = "D7DCE0"
DARK_GREY = "39434C"
AMBER = "C9841D"
RED = "A74444"
WHITE = "FFFFFF"

REQUIRED_TOP_LEVEL = {
    "meta",
    "executive_thesis",
    "executive_summary",
    "method_notes",
    "legal_structure",
    "orgchart",
    "power_map",
    "raci",
    "people",
    "influence_map",
    "contact_order",
    "unknowns",
    "sources",
}
REQUIRED_PERSON = {
    "name",
    "role",
    "perimeter",
    "experience",
    "expertise",
    "ai_communication",
    "confidence",
    "source_ids",
}
NODE_ID = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")
SOURCE_ID = re.compile(r"^S[0-9]{2,}$")


def _require(condition: bool, message: str, errors: list[str]) -> None:
    if not condition:
        errors.append(message)


def validate_config(cfg: dict[str, Any]) -> list[str]:
    """Return human-readable contract violations without mutating input."""
    errors: list[str] = []
    missing = sorted(REQUIRED_TOP_LEVEL - cfg.keys())
    _require(not missing, f"Missing top-level keys: {', '.join(missing)}", errors)
    if missing:
        return errors

    meta = cfg.get("meta", {})
    for key in ("company", "title", "subtitle", "date", "evidence_cutoff", "purpose", "disclaimer"):
        _require(bool(meta.get(key)), f"meta.{key} is required", errors)

    sources = cfg.get("sources", [])
    source_ids: list[str] = []
    for index, source in enumerate(sources):
        sid = source.get("id", "")
        _require(bool(SOURCE_ID.match(sid)), f"sources[{index}].id must match S00", errors)
        source_ids.append(sid)
        for key in ("source_type", "title", "publisher", "publication_or_access_date", "url", "supports_fields"):
            _require(key in source and source[key] not in (None, ""), f"sources[{index}].{key} is required", errors)
        parsed = urlparse(str(source.get("url", "")))
        _require(parsed.scheme in {"http", "https"} and bool(parsed.netloc), f"sources[{index}].url must be HTTP(S)", errors)
    _require(len(source_ids) == len(set(source_ids)), "Source IDs must be unique", errors)

    for index, person in enumerate(cfg.get("people", [])):
        missing_person = sorted(REQUIRED_PERSON - person.keys())
        _require(not missing_person, f"people[{index}] missing: {', '.join(missing_person)}", errors)
        for sid in person.get("source_ids", []):
            _require(sid in source_ids, f"people[{index}] references unknown source {sid}", errors)

    nodes = cfg.get("orgchart", {}).get("nodes", [])
    node_ids = [node.get("id", "") for node in nodes]
    for index, node_id in enumerate(node_ids):
        _require(bool(NODE_ID.match(node_id)), f"orgchart.nodes[{index}].id is not DOT-safe", errors)
    _require(len(node_ids) == len(set(node_ids)), "Org-chart node IDs must be unique", errors)
    for index, edge in enumerate(cfg.get("orgchart", {}).get("edges", [])):
        _require(isinstance(edge, list) and len(edge) == 3, f"orgchart.edges[{index}] must contain source, target, style", errors)
        if isinstance(edge, list) and len(edge) == 3:
            source, target, style = edge
            _require(source in node_ids and target in node_ids, f"orgchart.edges[{index}] references an unknown node", errors)
            _require(style in {"solid", "dashed"}, f"orgchart.edges[{index}] style must be solid or dashed", errors)

    headers = cfg.get("raci", {}).get("headers", [])
    for index, row in enumerate(cfg.get("raci", {}).get("rows", [])):
        _require(len(row) == len(headers), f"raci.rows[{index}] has {len(row)} cells; expected {len(headers)}", errors)
    return errors


def shade(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def border(cell: Any, color: str = MID_GREY, size: str = "5") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def cell_margins(cell: Any, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def no_split(row: Any) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def repeat_header(row: Any) -> None:
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(element)


def set_font(run: Any, size: float = 9, bold: bool = False, color: str = DARK_GREY, name: str = "Aptos", italic: bool = False) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def hyperlink(paragraph: Any, text: str, url: str) -> None:
    relation = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    link.append(run)
    paragraph._p.append(link)


def setup(document: Document, short_title: str, date: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = section.right_margin = Cm(1.7)
    section.header_distance = section.footer_distance = Cm(0.65)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(DARK_GREY)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in (
        ("Title", 26, NAVY, 0, 10),
        ("Heading 1", 17, NAVY, 12, 6),
        ("Heading 2", 13, DARK_GREEN, 10, 4),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display" if name != "Heading 2" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Source Note" not in document.styles:
        source_style = document.styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
        source_style.font.name = "Aptos"
        source_style.font.size = Pt(7.5)
        source_style.font.color.rgb = RGBColor.from_string("67727D")
    header = section.header.paragraphs[0]
    set_font(header.add_run(short_title), 7.5, True, NAVY)
    set_font(header.add_run(f"  |  Sources publiques — {date}"), 7.5, color="6C7780")
    footer = section.footer.add_table(rows=1, cols=2, width=Cm(17.6))
    set_font(footer.cell(0, 0).paragraphs[0].add_run("Analyse externe — organigramme reconstitué, non officiel"), 7, color="77828B")
    page_number(footer.cell(0, 1).paragraphs[0])


def callout(document: Document, title: str, body: str, fill: str = LIGHT_GREEN, line: str = GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    border(cell, line, "8")
    cell_margins(cell, 130, 160, 120, 160)
    set_font(cell.paragraphs[0].add_run(title), 9.5, True, NAVY)
    set_font(cell.add_paragraph().add_run(body), 9)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def matrix(document: Document, headers: list[Any], rows: Iterable[list[Any]], font: float = 7.8, head: str = NAVY) -> Any:
    rows = list(rows)
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, value in enumerate(headers):
        cell = table.cell(0, index)
        shade(cell, head)
        cell_margins(cell, 75, 80, 75, 80)
        set_font(cell.paragraphs[0].add_run(str(value)), font, True, WHITE)
    repeat_header(table.rows[0])
    for row_index, row in enumerate(rows):
        no_split(table.rows[row_index + 1])
        for column_index, value in enumerate(row):
            cell = table.cell(row_index + 1, column_index)
            shade(cell, WHITE if row_index % 2 else LIGHT_GREY)
            border(cell, MID_GREY, "4")
            cell_margins(cell, 70, 80, 70, 80)
            set_font(cell.paragraphs[0].add_run(str(value)), font)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def bullets(document: Document, items: Iterable[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        paragraph = document.add_paragraph(style=style)
        paragraph.paragraph_format.space_after = Pt(2.2)
        set_font(paragraph.add_run(str(item)), 9.1)


def make_orgchart(data: dict[str, Any], output: Path) -> None:
    dot_binary = shutil.which("dot")
    colors = {
        "governance": (LIGHT_GREEN, GREEN),
        "business": (LIGHT_BLUE, TEAL),
        "technology": ("E8F0FA", "315E8C"),
        "control": ("F0EAF7", "76539A"),
        "transition": ("FFF4E5", AMBER),
        "unknown": ("FCEAEA", RED),
        "federated": ("F3F6F8", "B9C4CC"),
    }
    lines = [
        "digraph G {",
        'graph [rankdir=TB, bgcolor="white", nodesep=0.28, ranksep=0.42, margin=0.05];',
        'node [shape=box, style="rounded,filled", fontname="Arial", fontsize=9, margin="0.10,0.07"];',
        'edge [color="#8A969F", arrowsize=0.55];',
    ]
    for node in data["nodes"]:
        fill, line = colors.get(node.get("tier", "federated"), colors["federated"])
        label = str(node["label"]).replace('"', r'\"')
        lines.append(f'{node["id"]} [label="{label}", fillcolor="#{fill}", color="#{line}"];')
    for source, target, style in data["edges"]:
        suffix = " [style=dashed]" if style == "dashed" else ""
        lines.append(f"{source} -> {target}{suffix};")
    lines.append("}")
    dot_path = output.with_suffix(".dot")
    dot_path.write_text("\n".join(lines), encoding="utf-8")
    if dot_binary:
        subprocess.run([dot_binary, "-Tpng", "-Gdpi=170", str(dot_path), "-o", str(output)], check=True)
    else:
        make_orgchart_fallback(data, output, colors)


def make_orgchart_fallback(data: dict[str, Any], output: Path, colors: dict[str, tuple[str, str]]) -> None:
    """Render a readable layered PNG when the Graphviz binary is unavailable."""
    nodes = {node["id"]: node for node in data["nodes"]}
    incoming = {node_id: 0 for node_id in nodes}
    children: dict[str, list[str]] = {node_id: [] for node_id in nodes}
    for source, target, _style in data["edges"]:
        incoming[target] += 1
        children[source].append(target)
    ranks = {node_id: 0 for node_id in nodes}
    queue = [node_id for node_id, count in incoming.items() if count == 0]
    visited: set[str] = set()
    while queue:
        source = queue.pop(0)
        visited.add(source)
        for target in children[source]:
            ranks[target] = max(ranks[target], ranks[source] + 1)
            incoming[target] -= 1
            if incoming[target] == 0:
                queue.append(target)
    for node_id in nodes.keys() - visited:
        ranks[node_id] = max(ranks.values(), default=0) + 1

    layers: dict[int, list[str]] = {}
    for node_id, rank in ranks.items():
        layers.setdefault(rank, []).append(node_id)
    box_width, box_height, x_gap, y_gap, margin = 300, 92, 34, 74, 40
    max_columns = max((len(layer) for layer in layers.values()), default=1)
    width = max(900, margin * 2 + max_columns * box_width + (max_columns - 1) * x_gap)
    height = margin * 2 + len(layers) * box_height + max(0, len(layers) - 1) * y_gap
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    bold_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    font = ImageFont.truetype(font_path, 17) if Path(font_path).exists() else ImageFont.load_default()
    bold = ImageFont.truetype(bold_path, 17) if Path(bold_path).exists() else font
    positions: dict[str, tuple[int, int, int, int]] = {}
    for rank in sorted(layers):
        layer = layers[rank]
        content_width = len(layer) * box_width + max(0, len(layer) - 1) * x_gap
        start_x = (width - content_width) // 2
        y = margin + rank * (box_height + y_gap)
        for index, node_id in enumerate(layer):
            x = start_x + index * (box_width + x_gap)
            positions[node_id] = (x, y, x + box_width, y + box_height)

    for source, target, style in data["edges"]:
        source_box, target_box = positions[source], positions[target]
        start = ((source_box[0] + source_box[2]) // 2, source_box[3])
        end = ((target_box[0] + target_box[2]) // 2, target_box[1])
        if style == "dashed":
            segments = 12
            for index in range(0, segments, 2):
                p1 = (start[0] + (end[0] - start[0]) * index // segments, start[1] + (end[1] - start[1]) * index // segments)
                p2 = (start[0] + (end[0] - start[0]) * (index + 1) // segments, start[1] + (end[1] - start[1]) * (index + 1) // segments)
                draw.line([p1, p2], fill="#8A969F", width=3)
        else:
            draw.line([start, end], fill="#8A969F", width=3)
        draw.polygon([(end[0] - 6, end[1] - 9), (end[0] + 6, end[1] - 9), end], fill="#8A969F")

    for node_id, node in nodes.items():
        box = positions[node_id]
        fill, line = colors.get(node.get("tier", "federated"), colors["federated"])
        draw.rounded_rectangle(box, radius=12, fill=f"#{fill}", outline=f"#{line}", width=3)
        raw_lines = str(node["label"]).replace("\\n", "\n").splitlines()
        lines = [wrapped for raw in raw_lines for wrapped in textwrap.wrap(raw, width=29) or [""]]
        line_height = 22
        text_y = box[1] + max(8, (box_height - len(lines) * line_height) // 2)
        for line_index, line_text in enumerate(lines[:4]):
            active_font = bold if line_index == 0 else font
            bounds = draw.textbbox((0, 0), line_text, font=active_font)
            text_x = box[0] + (box_width - (bounds[2] - bounds[0])) // 2
            draw.text((text_x, text_y), line_text, font=active_font, fill=f"#{DARK_GREY}")
            text_y += line_height
    image.save(output, "PNG", dpi=(170, 170))


def cover(document: Document, meta: dict[str, Any]) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(44)
    badge = document.add_paragraph()
    badge_shade = OxmlElement("w:shd")
    badge_shade.set(qn("w:fill"), GREEN)
    badge._p.get_or_add_pPr().append(badge_shade)
    set_font(badge.add_run(f"NOTE {meta.get('note_number', 2)}"), 10, True, WHITE)
    set_font(document.add_paragraph(style="Title").add_run(meta["title"]), 26, True, NAVY, "Aptos Display")
    set_font(document.add_paragraph().add_run(meta["subtitle"]), 13, True, DARK_GREEN)
    keywords = meta.get("keywords", [])
    if keywords:
        table = document.add_table(rows=1, cols=len(keywords))
        for index, keyword in enumerate(keywords):
            cell = table.cell(0, index)
            shade(cell, LIGHT_GREEN)
            border(cell, "B8D3AA")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(cell.paragraphs[0].add_run(keyword), 8.2, True, DARK_GREEN)
    paragraph = document.add_paragraph()
    set_font(paragraph.add_run("Objectif — "), 9, True, TEAL)
    set_font(paragraph.add_run(meta["purpose"]), 9)
    callout(document, "Statut du document", f"{meta['date']} — {meta['disclaimer']}", LIGHT_BLUE, TEAL)
    document.add_page_break()


def person_card(cell: Any, person: dict[str, Any]) -> None:
    shade(cell, WHITE)
    border(cell, GREEN, "8")
    cell_margins(cell, 110, 120, 105, 120)
    set_font(cell.paragraphs[0].add_run(person["name"]), 10, True, NAVY)
    set_font(cell.add_paragraph().add_run(person["role"]), 8.5, True, DARK_GREEN)
    fields = (
        ("Périmètre", "perimeter"),
        ("Mandat légal", "legal_mandate"),
        ("Ancienneté", "tenure"),
        ("Expérience", "experience"),
        ("Expertise", "expertise"),
        ("Communication IA", "ai_communication"),
        ("Confiance", "confidence"),
        ("Dernière vérification", "last_verified"),
    )
    for label, key in fields:
        if not person.get(key):
            continue
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.5)
        set_font(paragraph.add_run(f"{label} : "), 7.45, True, TEAL)
        set_font(paragraph.add_run(str(person[key])), 7.45)
    paragraph = cell.add_paragraph()
    set_font(paragraph.add_run("Sources : "), 7.45, True, TEAL)
    set_font(paragraph.add_run(", ".join(person["source_ids"])), 7.45)


def build(cfg: dict[str, Any], output: Path, keep_intermediates: bool = False) -> Path:
    errors = validate_config(cfg)
    if errors:
        raise ValueError("Invalid configuration:\n- " + "\n- ".join(errors))
    output = output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    work_context = tempfile.TemporaryDirectory(prefix="org-note-")
    workdir = Path(work_context.name)
    try:
        orgchart = workdir / "orgchart.png"
        make_orgchart(cfg["orgchart"], orgchart)
        document = Document()
        meta = cfg["meta"]
        setup(document, f"{meta['company']} — Organisation Tech", meta["date"])
        cover(document, meta)
        document.add_heading("Synthèse exécutive", 1)
        callout(document, "Conclusion", cfg["executive_thesis"])
        matrix(document, ["Point", "Lecture"], cfg["executive_summary"], 8.1)
        document.add_heading("Méthode et confiance", 1)
        matrix(
            document,
            ["Niveau", "Signification"],
            [
                ["Confirmé", "Source directe, officielle ou explicite et suffisamment récente."],
                ["Probable", "Plusieurs indices publics cohérents sans validation interne."],
                ["À confirmer", "Rôle, périmètre ou relation insuffisamment documenté."],
            ],
            8.1,
            TEAL,
        )
        bullets(document, cfg["method_notes"])
        document.add_heading("1. Structure juridique et portefeuille", 1)
        matrix(document, ["Niveau", "Rôle reconstitué", "Entités / exemples"], cfg["legal_structure"])
        document.add_heading("2. Organigramme de gouvernance et technologie", 1)
        document.add_picture(str(orgchart), width=Cm(16.9))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        note = document.add_paragraph(style="Source Note")
        note.add_run("Organigramme analytique reconstitué. Les pointillés indiquent une relation fonctionnelle ou probable.")
        document.add_heading("Lecture des pouvoirs", 2)
        matrix(document, ["Acteur", "Pouvoir probable", "Question de qualification"], cfg["power_map"])
        document.add_heading("3. RACI indicatif", 1)
        matrix(document, cfg["raci"]["headers"], cfg["raci"]["rows"], 7.5)
        document.add_paragraph("RACI hypothétique à valider.", style="Source Note")
        document.add_heading("4. Fiches décideurs et responsables", 1)
        callout(document, "Précaution", "Revalider le rôle et l’activité récente avant toute approche nominative.", LIGHT_BLUE, TEAL)
        cards = document.add_table(rows=0, cols=2)
        cards.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index in range(0, len(cfg["people"]), 2):
            row = cards.add_row()
            no_split(row)
            person_card(row.cells[0], cfg["people"][index])
            if index + 1 < len(cfg["people"]):
                person_card(row.cells[1], cfg["people"][index + 1])
            else:
                shade(row.cells[1], WHITE)
        document.add_heading("5. Cartographie d’influence", 1)
        matrix(document, ["Cercle", "Personnes / fonctions", "Rôle", "Angle"], cfg["influence_map"], 7.5)
        document.add_heading("Ordre de contact recommandé", 2)
        bullets(document, cfg["contact_order"], True)
        document.add_heading("6. Zones à confirmer", 1)
        bullets(document, cfg["unknowns"])
        document.add_page_break()
        document.add_heading("Sources et méthode", 1)
        callout(document, "Règle de lecture", "Mandats légaux, rôles publics et relations inférées restent des catégories distinctes.", LIGHT_BLUE, TEAL)
        for source in cfg["sources"]:
            paragraph = document.add_paragraph(style="Source Note")
            set_font(paragraph.add_run(f"{source['id']} — {source['title']}. "), 7.8, True, NAVY)
            set_font(paragraph.add_run(f"{source['publisher']} {source['publication_or_access_date']}. "), 7.8)
            hyperlink(paragraph, source["url"], source["url"])
        document.save(output)
        if keep_intermediates:
            shutil.copy2(orgchart, output.with_name(f"{output.stem}_orgchart.png"))
            shutil.copy2(orgchart.with_suffix(".dot"), output.with_name(f"{output.stem}_orgchart.dot"))
    finally:
        work_context.cleanup()
    return output


def load_json(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"Cannot read JSON config {path}: {exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("The configuration root must be an object")
    return data


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--config", type=Path, required=True)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--validate-only", action="store_true")
    parser.add_argument("--keep-intermediates", action="store_true")
    args = parser.parse_args()
    if not args.validate_only and args.output is None:
        parser.error("--output is required unless --validate-only is used")
    try:
        config = load_json(args.config)
        errors = validate_config(config)
        if errors:
            print("Configuration invalid:\n- " + "\n- ".join(errors), file=sys.stderr)
            return 2
        if args.validate_only:
            print(f"Configuration valid: {args.config}")
            return 0
        result = build(config, args.output, args.keep_intermediates)
        print(result)
        return 0
    except (ValueError, RuntimeError, subprocess.CalledProcessError) as exc:
        print(str(exc), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
